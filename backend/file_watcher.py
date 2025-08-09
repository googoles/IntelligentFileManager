"""
File monitoring and watching system for the Research File Manager MVP.

This module provides comprehensive file watching capabilities with automatic
content extraction and database integration.
"""

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler, FileSystemEvent
import os
import hashlib
import time
from datetime import datetime
from typing import Optional, Dict, List, Set
import logging
from pathlib import Path
import threading
import json
import mimetypes

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# OCR processing imports
try:
    from ocr_processor import get_ocr_processor, is_ocr_available
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    get_ocr_processor = None
    is_ocr_available = lambda: False

from database import db_manager, File

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class FileContentExtractor:
    """Handles content extraction from various file types including OCR for images and scanned PDFs."""
    
    # Supported text file extensions
    TEXT_EXTENSIONS = {'.txt', '.md', '.py', '.js', '.json', '.csv', '.xml', '.html', '.css', '.sql', '.r', '.java', '.cpp', '.c', '.h'}
    
    # Document file extensions that support text extraction
    DOCUMENT_EXTENSIONS = {'.pdf', '.docx', '.doc'}
    
    # Image file extensions that support OCR
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp'}
    
    # Maximum file sizes for processing (in bytes)
    MAX_PDF_SIZE = 50 * 1024 * 1024  # 50MB
    MAX_DOCX_SIZE = 20 * 1024 * 1024  # 20MB
    MAX_TEXT_SIZE = 10 * 1024 * 1024  # 10MB
    MAX_IMAGE_SIZE = 50 * 1024 * 1024  # 50MB for OCR
    
    # Maximum content length to extract (5000 characters as per MVP spec)
    MAX_CONTENT_LENGTH = 5000
    
    @classmethod
    def extract_content(cls, file_path: str) -> str:
        """
        Extract text content from supported file types including PDFs, Word documents, and images (via OCR).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content (up to MAX_CONTENT_LENGTH characters)
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            file_size = os.path.getsize(file_path)
            
            # Handle text files
            if file_ext in cls.TEXT_EXTENSIONS:
                if file_size > cls.MAX_TEXT_SIZE:
                    logger.warning(f"Text file too large for content extraction: {file_path}")
                    return ""
                
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read(cls.MAX_CONTENT_LENGTH)
                    logger.debug(f"Extracted {len(content)} characters from text file: {file_path}")
                    return content
            
            # Handle PDF files (try regular extraction first, then OCR for scanned PDFs)
            elif file_ext == '.pdf' and PDF_AVAILABLE:
                content = cls._extract_pdf_content(file_path, file_size)
                # If no text extracted from PDF, try OCR
                if not content.strip() and OCR_AVAILABLE:
                    logger.info(f"No text found in PDF, attempting OCR: {file_path}")
                    content = cls._extract_ocr_content(file_path, file_size)
                return content
            
            # Handle Word documents
            elif file_ext in {'.docx', '.doc'} and DOCX_AVAILABLE:
                return cls._extract_docx_content(file_path, file_size)
            
            # Handle image files with OCR
            elif file_ext in cls.IMAGE_EXTENSIONS and OCR_AVAILABLE:
                return cls._extract_ocr_content(file_path, file_size)
            
            else:
                logger.debug(f"Unsupported file type for content extraction: {file_ext}")
                return ""
                
        except (IOError, OSError, UnicodeDecodeError) as e:
            logger.warning(f"Failed to extract content from {file_path}: {e}")
            return ""
        except Exception as e:
            logger.error(f"Unexpected error extracting content from {file_path}: {e}")
            return ""
    
    @classmethod
    def get_file_metadata(cls, file_path: str) -> Dict:
        """
        Extract file metadata.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing file metadata
        """
        try:
            stat_info = os.stat(file_path)
            file_path_obj = Path(file_path)
            
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            file_hash = cls._calculate_file_hash(file_path)
            
            metadata = {
                'size': stat_info.st_size,
                'modified_time': stat_info.st_mtime,
                'created_time': getattr(stat_info, 'st_birthtime', stat_info.st_mtime),  # macOS only
                'extension': file_path_obj.suffix.lower(),
                'mime_type': mime_type,
                'is_hidden': file_path_obj.name.startswith('.'),
                'parent_directory': str(file_path_obj.parent),
                'file_hash': file_hash,
                'supports_text_extraction': file_path_obj.suffix.lower() in (cls.TEXT_EXTENSIONS | cls.DOCUMENT_EXTENSIONS),
                'supports_ocr': file_path_obj.suffix.lower() in cls.IMAGE_EXTENSIONS or (file_path_obj.suffix.lower() == '.pdf' and OCR_AVAILABLE),
                'processing_status': 'completed'
            }
            
            # Add document-specific metadata
            if file_path_obj.suffix.lower() == '.pdf':
                metadata['document_type'] = 'pdf'
                metadata['pdf_processing_available'] = PDF_AVAILABLE
            elif file_path_obj.suffix.lower() in {'.docx', '.doc'}:
                metadata['document_type'] = 'word'
                metadata['docx_processing_available'] = DOCX_AVAILABLE
            elif file_path_obj.suffix.lower() in cls.IMAGE_EXTENSIONS:
                metadata['document_type'] = 'image'
                metadata['ocr_available'] = OCR_AVAILABLE and is_ocr_available()
            
            return metadata
            
        except (IOError, OSError) as e:
            logger.warning(f"Failed to extract metadata from {file_path}: {e}")
            return {}
    
    @classmethod
    def _extract_pdf_content(cls, file_path: str, file_size: int) -> str:
        """
        Extract text content from PDF files using PyPDF2.
        
        Args:
            file_path: Path to the PDF file
            file_size: Size of the file in bytes
            
        Returns:
            Extracted text content
        """
        if not PDF_AVAILABLE:
            logger.debug(f"PyPDF2 not available for PDF extraction: {file_path}")
            return ""
        
        if file_size > cls.MAX_PDF_SIZE:
            logger.warning(f"PDF file too large for content extraction: {file_path}")
            return ""
        
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Limit to first 10 pages for performance
                max_pages = min(10, len(pdf_reader.pages))
                
                for page_num in range(max_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                            
                        # Stop if we have enough content
                        if len(''.join(text_content)) >= cls.MAX_CONTENT_LENGTH:
                            break
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num} from PDF {file_path}: {e}")
                        continue
                
                content = '\n'.join(text_content)[:cls.MAX_CONTENT_LENGTH]
                logger.debug(f"Extracted {len(content)} characters from PDF: {file_path}")
                return content
                
        except Exception as e:
            logger.warning(f"Failed to extract PDF content from {file_path}: {e}")
            return ""
    
    @classmethod
    def _extract_docx_content(cls, file_path: str, file_size: int) -> str:
        """
        Extract text content from Word documents using python-docx.
        
        Args:
            file_path: Path to the Word document
            file_size: Size of the file in bytes
            
        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            logger.debug(f"python-docx not available for Word extraction: {file_path}")
            return ""
        
        if file_size > cls.MAX_DOCX_SIZE:
            logger.warning(f"Word document too large for content extraction: {file_path}")
            return ""
        
        try:
            # Only process .docx files (not .doc which requires different library)
            if not file_path.lower().endswith('.docx'):
                logger.debug(f"Only .docx files supported, skipping: {file_path}")
                return ""
            
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    
                # Stop if we have enough content
                if len('\n'.join(text_content)) >= cls.MAX_CONTENT_LENGTH:
                    break
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
                        
                    # Stop if we have enough content
                    if len('\n'.join(text_content)) >= cls.MAX_CONTENT_LENGTH:
                        break
                        
                if len('\n'.join(text_content)) >= cls.MAX_CONTENT_LENGTH:
                    break
            
            content = '\n'.join(text_content)[:cls.MAX_CONTENT_LENGTH]
            logger.debug(f"Extracted {len(content)} characters from Word document: {file_path}")
            return content
            
        except Exception as e:
            logger.warning(f"Failed to extract Word document content from {file_path}: {e}")
            return ""
    
    @classmethod
    def _extract_ocr_content(cls, file_path: str, file_size: int) -> str:
        """
        Extract text content from images or scanned PDFs using OCR.
        
        Args:
            file_path: Path to the image or PDF file
            file_size: Size of the file in bytes
            
        Returns:
            Extracted text content via OCR
        """
        if not OCR_AVAILABLE or not is_ocr_available():
            logger.debug(f"OCR not available for content extraction: {file_path}")
            return ""
        
        file_ext = Path(file_path).suffix.lower()
        
        # Check size limits
        if file_ext in cls.IMAGE_EXTENSIONS and file_size > cls.MAX_IMAGE_SIZE:
            logger.warning(f"Image file too large for OCR: {file_path}")
            return ""
        elif file_ext == '.pdf' and file_size > cls.MAX_PDF_SIZE:
            logger.warning(f"PDF file too large for OCR: {file_path}")
            return ""
        
        try:
            # Get OCR processor
            ocr_processor = get_ocr_processor()
            if not ocr_processor:
                logger.warning(f"OCR processor not available: {file_path}")
                return ""
            
            # Check if file can be processed
            if not ocr_processor.can_process_file(file_path):
                logger.debug(f"File cannot be processed with OCR: {file_path}")
                return ""
            
            # Perform OCR extraction
            logger.info(f"Starting OCR extraction for: {file_path}")
            ocr_result = ocr_processor.extract_text(file_path)
            
            if ocr_result.error:
                logger.warning(f"OCR extraction failed for {file_path}: {ocr_result.error}")
                return ""
            
            if ocr_result.text:
                # Limit content length as per other extraction methods
                content = ocr_result.text[:cls.MAX_CONTENT_LENGTH]
                logger.info(f"OCR extracted {len(content)} characters from {file_path} "
                          f"(confidence: {ocr_result.confidence:.2f}, "
                          f"processing time: {ocr_result.processing_time:.2f}s)")
                return content
            else:
                logger.debug(f"No text extracted via OCR from {file_path}")
                return ""
                
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {e}")
            return ""
    
    @staticmethod
    def _calculate_file_hash(file_path: str, chunk_size: int = 8192) -> Optional[str]:
        """Calculate MD5 hash of file for duplicate detection."""
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(chunk_size), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.warning(f"Failed to calculate hash for {file_path}: {e}")
            return None


class FileHandler(FileSystemEventHandler):
    """
    File system event handler for monitoring project directories.
    
    Handles file creation, modification, and deletion events with
    automatic content extraction and database updates.
    """
    
    def __init__(self, project_id: str, project_path: str):
        """
        Initialize file handler.
        
        Args:
            project_id: UUID of the project
            project_path: Root path of the project to monitor
        """
        super().__init__()
        self.project_id = project_id
        self.project_path = project_path
        self.content_extractor = FileContentExtractor()
        
        # Set to track recently processed files (prevents duplicate processing)
        self._recently_processed: Set[str] = set()
        self._processing_lock = threading.Lock()
        
        # Cleanup recently processed files every 5 minutes
        self._cleanup_timer = threading.Timer(300.0, self._cleanup_recently_processed)
        self._cleanup_timer.daemon = True
        self._cleanup_timer.start()
        
        logger.info(f"FileHandler initialized for project {project_id} at {project_path}")
    
    def on_created(self, event: FileSystemEvent) -> None:
        """Handle file creation events."""
        if not event.is_directory:
            self._process_file_event(event.src_path, event_type="created")
    
    def on_modified(self, event: FileSystemEvent) -> None:
        """Handle file modification events."""
        if not event.is_directory:
            self._process_file_event(event.src_path, event_type="modified")
    
    def on_deleted(self, event: FileSystemEvent) -> None:
        """Handle file deletion events."""
        if not event.is_directory:
            self._handle_file_deletion(event.src_path)
    
    def on_moved(self, event: FileSystemEvent) -> None:
        """Handle file move/rename events."""
        if not event.is_directory:
            # Handle as deletion of old path and creation of new path
            self._handle_file_deletion(event.src_path)
            self._process_file_event(event.dest_path, event_type="moved")
    
    def _process_file_event(self, file_path: str, event_type: str) -> None:
        """
        Process a file system event.
        
        Args:
            file_path: Path to the affected file
            event_type: Type of event (created, modified, moved)
        """
        try:
            # Normalize path
            file_path = os.path.abspath(file_path)
            
            # Skip if file doesn't exist (race condition)
            if not os.path.exists(file_path):
                logger.debug(f"File no longer exists, skipping: {file_path}")
                return
            
            # Skip temporary files and system files
            if self._should_ignore_file(file_path):
                logger.debug(f"Ignoring file: {file_path}")
                return
            
            # Prevent duplicate processing
            with self._processing_lock:
                if file_path in self._recently_processed:
                    logger.debug(f"File recently processed, skipping: {file_path}")
                    return
                self._recently_processed.add(file_path)
            
            # Process the file
            self.process_file(file_path, event_type)
            
        except Exception as e:
            logger.error(f"Error processing file event for {file_path}: {e}")
    
    def process_file(self, file_path: str, event_type: str = "created") -> Optional[File]:
        """
        Process a file and add it to the database.
        
        Args:
            file_path: Path to the file to process
            event_type: Type of event that triggered processing
            
        Returns:
            File object if successfully processed, None otherwise
        """
        try:
            # Ensure file exists and is accessible
            if not os.path.exists(file_path):
                logger.warning(f"File does not exist: {file_path}")
                return None
            
            # Extract basic file information
            file_name = os.path.basename(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Extract metadata
            metadata = self.content_extractor.get_file_metadata(file_path)
            metadata['event_type'] = event_type
            metadata['processed_at'] = datetime.utcnow().isoformat()
            
            # Extract content
            content = self.content_extractor.extract_content(file_path)
            
            # Add OCR metadata if applicable
            if file_ext.lower() in self.content_extractor.IMAGE_EXTENSIONS or \
               (file_ext.lower() == '.pdf' and OCR_AVAILABLE):
                try:
                    ocr_processor = get_ocr_processor()
                    if ocr_processor and ocr_processor.can_process_file(file_path):
                        ocr_status = ocr_processor.get_file_ocr_status(file_path)
                        metadata.update({
                            'ocr_status': ocr_status,
                            'ocr_processed': bool(content),  # True if content was extracted
                            'content_source': 'ocr' if file_ext.lower() in self.content_extractor.IMAGE_EXTENSIONS else 'mixed'
                        })
                except Exception as e:
                    logger.warning(f"Failed to get OCR metadata for {file_path}: {e}")
            
            # Add to database
            file_obj = db_manager.add_file(
                project_id=self.project_id,
                file_path=file_path,
                name=file_name,
                file_type=file_ext,
                content=content,
                metadata=metadata
            )
            
            logger.info(f"✅ File processed: {file_path} ({event_type})")
            return file_obj
            
        except Exception as e:
            logger.error(f"Failed to process file {file_path}: {e}")
            return None
    
    def _handle_file_deletion(self, file_path: str) -> None:
        """Handle file deletion events."""
        try:
            file_path = os.path.abspath(file_path)
            success = db_manager.delete_file(file_path)
            if success:
                logger.info(f"🗑️ File deleted from database: {file_path}")
            else:
                logger.debug(f"File not found in database for deletion: {file_path}")
        except Exception as e:
            logger.error(f"Error handling file deletion for {file_path}: {e}")
    
    def _should_ignore_file(self, file_path: str) -> bool:
        """
        Determine if a file should be ignored.
        
        Args:
            file_path: Path to check
            
        Returns:
            True if file should be ignored, False otherwise
        """
        file_path_obj = Path(file_path)
        file_name = file_path_obj.name
        
        # Ignore patterns
        ignore_patterns = [
            # System files
            '.DS_Store', 'Thumbs.db', 'desktop.ini',
            # Temporary files
            '.tmp', '.temp', '~$',
            # IDE files
            '.vscode', '.idea', '__pycache__',
            # Version control
            '.git', '.svn', '.hg',
            # Build artifacts
            'node_modules', 'dist', 'build', 'target'
        ]
        
        # Check if file name matches ignore patterns
        for pattern in ignore_patterns:
            if pattern in file_name or file_name.startswith(pattern) or file_name.endswith(pattern):
                return True
        
        # Ignore very small files (likely empty or system files)
        try:
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                return True
            
            # Also ignore extremely large files that would be impractical to process
            max_size = 500 * 1024 * 1024  # 500MB limit
            if file_size > max_size:
                logger.info(f"Ignoring very large file ({file_size / (1024*1024):.1f}MB): {file_path}")
                return True
                
        except OSError:
            return True
        
        # Ignore files in ignore directories
        ignore_dirs = {'.git', '__pycache__', '.vscode', '.idea', 'node_modules', '.svn', '.hg'}
        for part in file_path_obj.parts:
            if part in ignore_dirs:
                return True
        
        return False
    
    def _cleanup_recently_processed(self) -> None:
        """Clean up the recently processed files set."""
        with self._processing_lock:
            self._recently_processed.clear()
        
        # Restart timer
        self._cleanup_timer = threading.Timer(300.0, self._cleanup_recently_processed)
        self._cleanup_timer.daemon = True
        self._cleanup_timer.start()
        
        logger.debug("Cleaned up recently processed files set")


class ProjectWatcher:
    """
    Manages file watching for a single project.
    """
    
    def __init__(self, project_id: str, project_path: str):
        """
        Initialize project watcher.
        
        Args:
            project_id: UUID of the project
            project_path: Root path of the project
        """
        self.project_id = project_id
        self.project_path = project_path
        self.observer = Observer()
        self.file_handler = FileHandler(project_id, project_path)
        self.is_watching = False
        
        logger.info(f"ProjectWatcher initialized for {project_id} at {project_path}")
    
    def start_watching(self) -> bool:
        """
        Start watching the project directory.
        
        Returns:
            True if watching started successfully, False otherwise
        """
        try:
            if self.is_watching:
                logger.warning(f"Already watching project {self.project_id}")
                return True
            
            # Ensure directory exists
            if not os.path.exists(self.project_path):
                logger.error(f"Project path does not exist: {self.project_path}")
                return False
            
            # Schedule watching with recursive monitoring
            self.observer.schedule(
                self.file_handler,
                self.project_path,
                recursive=True
            )
            
            # Start the observer
            self.observer.start()
            self.is_watching = True
            
            logger.info(f"🔍 Started watching project {self.project_id} at {self.project_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to start watching project {self.project_id}: {e}")
            return False
    
    def stop_watching(self) -> None:
        """Stop watching the project directory."""
        try:
            if not self.is_watching:
                return
            
            self.observer.stop()
            self.observer.join(timeout=5.0)  # Wait up to 5 seconds for clean shutdown
            self.is_watching = False
            
            logger.info(f"⏹️ Stopped watching project {self.project_id}")
            
        except Exception as e:
            logger.error(f"Error stopping watcher for project {self.project_id}: {e}")
    
    def index_existing_files(self) -> int:
        """
        Index all existing files in the project directory.
        
        Returns:
            Number of files successfully indexed
        """
        indexed_count = 0
        
        try:
            for root, dirs, files in os.walk(self.project_path):
                # Filter out ignored directories
                dirs[:] = [d for d in dirs if not self.file_handler._should_ignore_file(os.path.join(root, d))]
                
                for file_name in files:
                    file_path = os.path.join(root, file_name)
                    
                    if not self.file_handler._should_ignore_file(file_path):
                        file_obj = self.file_handler.process_file(file_path, "indexed")
                        if file_obj:
                            indexed_count += 1
            
            logger.info(f"📚 Indexed {indexed_count} existing files for project {self.project_id}")
            return indexed_count
            
        except Exception as e:
            logger.error(f"Error indexing existing files for project {self.project_id}: {e}")
            return indexed_count


class FileWatcherManager:
    """
    Global manager for all project watchers.
    """
    
    def __init__(self):
        """Initialize the file watcher manager."""
        self.watchers: Dict[str, ProjectWatcher] = {}
        self._lock = threading.Lock()
        logger.info("FileWatcherManager initialized")
    
    def start_watching_project(self, project_id: str, project_path: str, index_existing: bool = True) -> bool:
        """
        Start watching a project directory.
        
        Args:
            project_id: UUID of the project
            project_path: Root path of the project
            index_existing: Whether to index existing files
            
        Returns:
            True if watching started successfully, False otherwise
        """
        with self._lock:
            if project_id in self.watchers:
                logger.warning(f"Project {project_id} is already being watched")
                return True
            
            try:
                watcher = ProjectWatcher(project_id, project_path)
                
                if not watcher.start_watching():
                    return False
                
                self.watchers[project_id] = watcher
                
                # Index existing files in background if requested
                if index_existing:
                    threading.Thread(
                        target=watcher.index_existing_files,
                        daemon=True
                    ).start()
                
                logger.info(f"✅ Started watching project {project_id}")
                return True
                
            except Exception as e:
                logger.error(f"Failed to start watching project {project_id}: {e}")
                return False
    
    def stop_watching_project(self, project_id: str) -> None:
        """
        Stop watching a project directory.
        
        Args:
            project_id: UUID of the project
        """
        with self._lock:
            if project_id in self.watchers:
                self.watchers[project_id].stop_watching()
                del self.watchers[project_id]
                logger.info(f"Stopped watching project {project_id}")
    
    def stop_all(self) -> None:
        """Stop watching all projects."""
        with self._lock:
            for project_id in list(self.watchers.keys()):
                self.stop_watching_project(project_id)
        logger.info("Stopped all project watchers")
    
    def get_watching_projects(self) -> List[str]:
        """Get list of currently watched project IDs."""
        with self._lock:
            return list(self.watchers.keys())


# Global file watcher manager instance
file_watcher_manager = FileWatcherManager()


def start_watching(project_path: str, project_id: str, index_existing: bool = True) -> bool:
    """
    Convenience function to start watching a project directory.
    
    Args:
        project_path: Path to the project directory
        project_id: UUID of the project
        index_existing: Whether to index existing files
        
    Returns:
        True if watching started successfully, False otherwise
    """
    return file_watcher_manager.start_watching_project(project_id, project_path, index_existing)


def stop_watching(project_id: str) -> None:
    """
    Convenience function to stop watching a project directory.
    
    Args:
        project_id: UUID of the project
    """
    file_watcher_manager.stop_watching_project(project_id)


def stop_all_watchers() -> None:
    """Stop all file watchers."""
    file_watcher_manager.stop_all()